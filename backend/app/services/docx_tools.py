"""Word (.docx) export for Levy artifacts.

Users (especially self-representing litigants and lawyers) often need the
editable Word version of a draft, not just the PDF ("turn these documents to
word, I cannot convert"). We render Word lazily: every generated artifact
stores its source Markdown in `meta.source_markdown`, and the first time the
user asks for the Word version we render + cache it in the artifacts bucket.

Rendering path: Markdown -> HTML (reusing the same `markdown` lib the PDF path
uses, so court captions emitted as raw HTML pass straight through) -> a small
HTML walker that builds the python-docx document. Headings, bold/italic, lists,
block-quotes, horizontal rules and tables are supported; unknown tags degrade
to their text. python-docx (1.1.2) is already a dependency.
"""
from __future__ import annotations

import io
from html.parser import HTMLParser
from html import unescape

import httpx
import markdown as md_lib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from ..config import get_settings
from ..db.supabase import get_db


# ─── low-level docx helpers ──────────────────────────────────────────────────


def _bottom_border(paragraph) -> None:
    """Give a paragraph a thin bottom rule (used to render <hr>)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    pPr.append(borders)


class _HtmlToDocx(HTMLParser):
    """Walk the rendered HTML and emit python-docx blocks."""

    BLOCK_HEADINGS = {"h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self, doc: Document, *, legal: bool) -> None:
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.legal = legal
        self.para = None
        self.inline: list[str] = []          # active inline styles: b / i / code
        self.list_stack: list[str] = []        # ul / ol nesting
        self.align: int | None = None
        # table buffering (python-docx needs col count up front)
        self.in_table = False
        self.rows: list[list[str]] = []
        self.cur_row: list[str] | None = None
        self.cur_cell: list[str] | None = None
        self.cell_header = False

    # -- helpers --
    def _flush_para(self) -> None:
        self.para = None
        self.align = None

    def _ensure_para(self):
        if self.para is None:
            self.para = self.doc.add_paragraph()
            if self.align is not None:
                self.para.alignment = self.align
        return self.para

    def _emit_text(self, text: str) -> None:
        if self.in_table and self.cur_cell is not None:
            self.cur_cell.append(text)
            return
        if not text.strip() and self.para is None:
            return
        para = self._ensure_para()
        run = para.add_run(text)
        if "b" in self.inline:
            run.bold = True
        if "i" in self.inline:
            run.italic = True
        if "code" in self.inline:
            run.font.name = "Consolas"

    # -- parser callbacks --
    def handle_starttag(self, tag, attrs):  # noqa: C901
        a = dict(attrs)
        if tag in self.BLOCK_HEADINGS:
            level = min(int(tag[1]), 4)
            self.para = self.doc.add_heading("", level=level)
            if self.legal:
                self.para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == "p":
            self.para = self.doc.add_paragraph()
            if self.legal:
                self.para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif tag in ("strong", "b"):
            self.inline.append("b")
        elif tag in ("em", "i"):
            self.inline.append("i")
        elif tag == "code":
            self.inline.append("code")
        elif tag == "br":
            if self.para is not None:
                self.para.add_run().add_break()
        elif tag == "hr":
            _bottom_border(self.doc.add_paragraph())
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
        elif tag == "li":
            style = "List Number" if (self.list_stack and self.list_stack[-1] == "ol") else "List Bullet"
            try:
                self.para = self.doc.add_paragraph(style=style)
            except KeyError:
                self.para = self.doc.add_paragraph()
        elif tag == "blockquote":
            try:
                self.para = self.doc.add_paragraph(style="Quote")
            except KeyError:
                self.para = self.doc.add_paragraph()
                self.para.paragraph_format.left_indent = Pt(18)
        elif tag == "table":
            self.in_table = True
            self.rows = []
        elif tag == "tr" and self.in_table:
            self.cur_row = []
        elif tag in ("td", "th") and self.in_table:
            self.cur_cell = []
            self.cell_header = tag == "th"
        elif tag == "div":
            self.para = self.doc.add_paragraph()
            style = a.get("style", "")
            if "center" in style:
                self.para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.align = WD_ALIGN_PARAGRAPH.CENTER

    def handle_endtag(self, tag):
        if tag in self.BLOCK_HEADINGS or tag in ("p", "li", "blockquote", "div"):
            self._flush_para()
        elif tag in ("strong", "b") and "b" in self.inline:
            self.inline.remove("b")
        elif tag in ("em", "i") and "i" in self.inline:
            self.inline.remove("i")
        elif tag == "code" and "code" in self.inline:
            self.inline.remove("code")
        elif tag in ("ul", "ol") and self.list_stack:
            self.list_stack.pop()
        elif tag in ("td", "th") and self.in_table and self.cur_row is not None:
            self.cur_row.append("".join(self.cur_cell or []).strip())
            self.cur_cell = None
        elif tag == "tr" and self.in_table and self.cur_row is not None:
            self.rows.append(self.cur_row)
            self.cur_row = None
        elif tag == "table" and self.in_table:
            self._build_table()
            self.in_table = False
            self.rows = []

    def handle_data(self, data):
        if data:
            self._emit_text(unescape(data) if "&" in data else data)

    def _build_table(self) -> None:
        rows = [r for r in self.rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=0, cols=cols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        for r in rows:
            cells = table.add_row().cells
            for i in range(cols):
                cells[i].text = r[i] if i < len(r) else ""
        self._flush_para()


# ─── plain-text extraction (for "copy the draft" / mobile) ───────────────────


class _HtmlToText(HTMLParser):
    """Walk rendered HTML and emit clean, copyable plain text.

    Legal drafts store their source Markdown with raw HTML in it (court captions,
    parties tables), so returning that Markdown verbatim pastes tag soup. This
    flattens the SAME rendered HTML the PDF uses into readable lines: block
    elements break lines, <br> breaks, list items get a marker, table rows become
    tab-separated lines. Entities are unescaped (convert_charrefs=True).
    """

    _LINE_BLOCKS = {"p", "div", "blockquote", "li", "tr",
                    "h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.list_stack: list[str] = []
        self.ol_counts: list[int] = []
        self._first_cell = True

    def handle_starttag(self, tag, attrs):
        if tag == "br":
            self.parts.append("\n")
        elif tag == "hr":
            self.parts.append("\n" + "-" * 24 + "\n")
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
            self.ol_counts.append(0)
        elif tag == "li":
            self.parts.append("\n")
            if self.list_stack and self.list_stack[-1] == "ol":
                self.ol_counts[-1] += 1
                self.parts.append(f"{self.ol_counts[-1]}. ")
            else:
                self.parts.append("- ")
        elif tag == "tr":
            self.parts.append("\n")
            self._first_cell = True
        elif tag in ("td", "th"):
            if not self._first_cell:
                self.parts.append("\t")
            self._first_cell = False
        elif tag in self._LINE_BLOCKS:
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag == "ul" and self.list_stack:
            self.list_stack.pop()
        elif tag == "ol" and self.list_stack:
            self.list_stack.pop()
            if self.ol_counts:
                self.ol_counts.pop()
        elif tag in self._LINE_BLOCKS:
            self.parts.append("\n")

    def handle_data(self, data):
        if data:
            self.parts.append(data)


def markdown_to_plaintext(body_md: str) -> str:
    """Render a document's source Markdown (which may contain raw HTML) to clean
    plain text suitable for copying into an email or Word document."""
    html = md_lib.markdown(body_md or "", extensions=["extra", "tables", "sane_lists"])
    parser = _HtmlToText()
    parser.feed(html)
    parser.close()
    raw = "".join(parser.parts)
    # Collapse runs of blank lines to at most one, trim trailing spaces per line.
    out: list[str] = []
    blanks = 0
    for line in raw.split("\n"):
        line = line.rstrip()
        if not line.strip():
            blanks += 1
            if blanks <= 1:
                out.append("")
        else:
            blanks = 0
            out.append(line)
    return "\n".join(out).strip()


# ─── public render ───────────────────────────────────────────────────────────


def render_docx_bytes(*, title: str, body_md: str, subtitle: str | None = None, legal: bool = False) -> bytes:
    """Render a Markdown body to .docx bytes."""
    doc = Document()
    normal = doc.styles["Normal"].font
    if legal:
        normal.name = "Times New Roman"
        normal.size = Pt(12)
    else:
        normal.name = "Calibri"
        normal.size = Pt(11)

    if not legal:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0F, 0x3A, 0x2A)
        if subtitle:
            sub = doc.add_paragraph(subtitle)
            sub.runs[0].italic = True
            sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    html = md_lib.markdown(body_md, extensions=["extra", "tables", "sane_lists"])
    parser = _HtmlToDocx(doc, legal=legal)
    parser.feed(html)
    parser.close()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── storage + lazy artifact export ──────────────────────────────────────────

_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _upload_artifact_docx(artifact_id: str, data: bytes) -> str:
    settings = get_settings()
    base = settings.supabase_url.rstrip("/")
    key = f"{artifact_id}.docx"
    with httpx.Client(timeout=60.0) as client:
        client.delete(
            f"{base}/storage/v1/object/artifacts/{key}",
            headers={"Authorization": f"Bearer {settings.supabase_key}"},
        )
        resp = client.post(
            f"{base}/storage/v1/object/artifacts/{key}",
            headers={
                "Authorization": f"Bearer {settings.supabase_key}",
                "Content-Type": _DOCX_CT,
                "x-upsert": "true",
            },
            content=data,
        )
        resp.raise_for_status()
    return f"artifacts/{key}"


def ensure_artifact_docx(artifact_id: str) -> dict:
    """Return {storage_path, title} for the Word version, rendering it once.

    Requires the artifact to have `meta.source_markdown` (every Levy-generated
    document stores it). Raises ValueError if the artifact cannot be exported.
    """
    db = get_db()
    res = (
        db.table("artifacts")
        .select("id,title,meta,owner_id")
        .eq("id", artifact_id)
        .limit(1)
        .execute()
    )
    if not res.data:
        raise ValueError("artifact not found")
    row = res.data[0]
    meta = row.get("meta") or {}

    cached = meta.get("docx_path")
    if cached:
        return {"storage_path": cached, "title": row.get("title")}

    source_md = meta.get("source_markdown")
    if not source_md:
        raise ValueError("this document has no editable source to export to Word")

    data = render_docx_bytes(
        title=row.get("title") or "Document",
        body_md=source_md,
        subtitle=meta.get("subtitle"),
        legal=(meta.get("layout") == "legal"),
    )
    storage_path = _upload_artifact_docx(artifact_id, data)
    meta = {**meta, "docx_path": storage_path}
    db.table("artifacts").update({"meta": meta}).eq("id", artifact_id).execute()
    return {"storage_path": storage_path, "title": row.get("title")}
