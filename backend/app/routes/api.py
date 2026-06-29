"""
API Routes — FastAPI endpoints for Levy.

Three endpoints for the baseline RAG system:
  POST /api/chat       — Full RAG: retrieve + generate answer with citations
  POST /api/search     — Retrieval only: test what chunks come back (no LLM)
  GET  /api/documents  — List ingested documents and stats
"""

import os
import io
import re
import asyncio
import logging
import tempfile
import uuid
from datetime import datetime, timezone
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends, Header, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import json
import time
from ..services import rag
from ..services.agent import run_agent
from ..services.chat_persist import RunAccumulator
from ..services.ingester import ingest_pdf
from ..services.embedder import get_query_embedding
from ..db.supabase import search_chunks, get_db
from ..prompts.legal_qa import SYSTEM_PROMPT, build_context_prompt
from ..prompts.irac_brief import IRAC_SYSTEM_PROMPT, build_irac_prompt
from ..providers.anthropic_provider import generate_response, generate_response_stream
from ..models.schemas import BriefRequest
from ..config import get_settings
from ..auth import require_user, optional_user, verify_token

router = APIRouter(prefix="/api")
logger = logging.getLogger("levy.api")

# Strong refs to detached durable-run tasks so the event loop doesn't GC them
# mid-flight; each removes itself on completion.
_INFLIGHT_RUNS: set = set()


# --- Ownership guards (backend uses service_role, which bypasses RLS, so we
#     must authorize every owner-scoped row here in code) -------------------

def _assert_owns(table: str, row_id: str, uid: str, *, owner_col: str = "owner_id") -> dict:
    """Fetch a row and assert it belongs to `uid`. 404 if missing, 403 if not theirs."""
    res = get_db().table(table).select("*").eq("id", row_id).limit(1).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail=f"{table[:-1]} not found")
    row = res.data[0]
    if row.get(owner_col) != uid:
        raise HTTPException(status_code=403, detail="not authorized for this resource")
    return row


def _assert_owns_session(session_id: str, uid: str) -> dict:
    res = get_db().table("chat_sessions").select("id,user_id").eq("id", session_id).limit(1).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="session not found")
    if res.data[0].get("user_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this session")
    return res.data[0]


# --- Request/Response Models ---

class ChatRequest(BaseModel):
    query: str
    model: str | None = None
    top_k: int | None = None
    threshold: float | None = None
    web_search: bool = False
    history: list[dict] | None = None
    user_id: str | None = None
    session_id: str | None = None
    attached_doc_ids: list[str] | None = None


class SearchRequest(BaseModel):
    query: str
    top_k: int | None = None
    threshold: float | None = None


# --- Endpoints ---

@router.post("/chat")
def chat(request: ChatRequest):
    """
    Full RAG pipeline: embed query → search → generate answer with citations.

    This is the main endpoint users interact with.
    """
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    try:
        result = rag.query(
            question=request.query,
            model=request.model,
            top_k=request.top_k,
            threshold=request.threshold,
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail="request could not be completed")


@router.post("/search")
def search(request: SearchRequest):
    """
    Retrieval only — returns matching chunks without LLM generation.

    Use this endpoint to:
    - Test retrieval quality independently
    - Debug which chunks match a query
    - Evaluate Recall@K and Precision@K
    - Save LLM costs during development
    """
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    try:
        result = rag.search_only(
            question=request.query,
            top_k=request.top_k,
            threshold=request.threshold,
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail="request could not be completed")


# Abuse guard for the public, expensive, anonymous chat endpoint. Search
# crawlers (and scripts) hitting /chat?q=... links were triggering billed agent
# runs at high volume. We block automated user-agents outright and rate-limit
# anonymous callers per client IP. Authenticated users are unaffected.
_BOT_UA = re.compile(
    r"(bot|crawler|spider|slurp|crawl|headless|preview|python-requests|httpx|aiohttp|"
    r"curl|wget|scrapy|node-fetch|axios|go-http|java/|okhttp|libwww|phantom|puppeteer|playwright)",
    re.I,
)
_ANON_HITS: dict[str, list[float]] = {}
_ANON_LIMIT = 6        # max anonymous chat requests
_ANON_WINDOW = 60.0    # per 60 seconds, per client IP


def _client_ip(req: Request) -> str:
    xff = (req.headers.get("x-forwarded-for") or "").split(",")[0].strip()
    return xff or (req.client.host if req.client else "unknown")


@router.post("/chat/stream")
async def chat_stream(request: ChatRequest, http_request: Request, authorization: str | None = Header(default=None)):
    """
    Agentic RAG pipeline. Server-Sent Events.

    Drives the model through a tool-use loop with `search_corpus` always
    available and Tavily-backed web tools added when `web_search=True`.

    Event types streamed (each as `data: {json}\n\n`):
      thinking, tool_call, tool_result, token, sources, done, error.

    Backwards-compat: `token` and `sources` events keep the same shape the
    pre-agent client already understands; new clients can additionally render
    `tool_call`/`tool_result`.
    """
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    # Identity comes from the verified bearer token, NOT the client-supplied
    # user_id. Anonymous callers (no token) get global-corpus-only answers.
    token = (authorization or "")[7:].strip() if (authorization or "").lower().startswith("bearer ") else None
    uid = verify_token(token) if token else None

    # Abuse guard: block automated user-agents, and rate-limit anonymous callers
    # per IP. This stops crawlers/scripts from draining the LLM budget via the
    # public chat endpoint. Signed-in users are exempt.
    ua = http_request.headers.get("user-agent") or ""
    if _BOT_UA.search(ua):
        raise HTTPException(status_code=403, detail="Automated access to chat is not allowed.")
    if not uid:
        # TEMPORARY HARD STOP: anonymous chat is disabled. A distributed flood
        # (browser-like user-agents across many IPs) bypassed the per-IP rate
        # limit and kept draining the LLM budget. Require sign-in until we add a
        # bot challenge (e.g. Cloudflare Turnstile) to reopen anonymous access.
        raise HTTPException(status_code=401, detail="Please sign in to use Levy chat.")

    # Only honour a session_id the caller actually owns; otherwise ignore it
    # so nobody can read another user's attached documents via the agent.
    safe_session_id = None
    if request.session_id and uid:
        s = get_db().table("chat_sessions").select("user_id").eq("id", request.session_id).limit(1).execute()
        if s.data and s.data[0].get("user_id") == uid:
            safe_session_id = request.session_id

    def _sse(event: dict) -> str:
        # The pre-agent client expects `chunks_used` on the sources event.
        if event.get("type") == "sources":
            payload = {
                "type": "sources",
                "db": event.get("db", []),
                "web": event.get("web", []),
                "chunks_used": [
                    {
                        "id": s.get("id"), "act_name": s.get("act_name"),
                        "section": s.get("section"), "part": s.get("part"),
                        "page_start": s.get("page_start"), "page_end": s.get("page_end"),
                        "similarity": s.get("similarity"), "content_preview": s.get("content_preview"),
                    }
                    for s in event.get("db", [])[:8]
                ],
            }
            return f"data: {json.dumps(payload)}\n\n"
        return f"data: {json.dumps(event)}\n\n"

    # Tier-1 durable execution: drive the agent in a DETACHED task that pushes
    # events to a queue and, on completion, persists the assistant message
    # server-side. If the client disconnects (closed tab / lost mobile signal)
    # the response generator is cancelled, but this task keeps running and still
    # saves the reply — so the thread is never left with a dangling no-reply.
    queue: asyncio.Queue = asyncio.Queue()
    acc = RunAccumulator()

    async def _drive():
        try:
            async for event in run_agent(
                user_query=request.query,
                model=request.model,
                web_enabled=bool(request.web_search),
                history=request.history,
                owner_id=uid,
                session_id=safe_session_id,
                attached_doc_ids=request.attached_doc_ids,
            ):
                acc.consume(event)
                await queue.put(event)
        except Exception as e:  # noqa: BLE001
            logger.exception("agent run failed")
            msg = (
                "Levy ran into a problem answering that. Please try again. If you pasted a very "
                "long document or table, try sending it in smaller parts."
            )
            # Persist a visible message so a failed run is never saved as a blank
            # reply (which made users re-send the same large paste repeatedly).
            if not acc.has_content():
                acc.content = msg
                acc.blocks.append({"kind": "text", "text": msg})
            await queue.put({"type": "error", "message": msg})
        finally:
            await queue.put(None)  # stream sentinel
            # Durable, server-owned save (signed-in threads only). Anonymous
            # callers have no session and aren't persisted.
            if safe_session_id and uid:
                try:
                    await asyncio.to_thread(acc.save, safe_session_id)
                except Exception:
                    logger.exception("durable save failed")

    task = asyncio.create_task(_drive())
    _INFLIGHT_RUNS.add(task)
    task.add_done_callback(_INFLIGHT_RUNS.discard)

    async def event_stream():
        while True:
            event = await queue.get()
            if event is None:
                break
            yield _sse(event)
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@router.get("/documents/{document_id}/pdf")
def get_document_pdf_url(document_id: str, expires_in: int = 3600, uid: str | None = Depends(optional_user)):
    """
    Return a short-lived signed URL for the canonical PDF of a legal document.

    Global-library docs are readable by anyone; user-uploaded docs only by
    their owner. The PDF lives in the private `legal-docs` bucket; we mint a
    signed URL on demand for the PDF.js viewer.
    """
    from ..db.supabase import get_db

    db = get_db()
    res = (
        db.table("legal_documents")
        .select("id, title, short_name, pdf_storage_path, pdf_page_count, canonical_url, is_global, owner_id")
        .eq("id", document_id)
        .limit(1)
        .execute()
    )
    if not res.data:
        raise HTTPException(status_code=404, detail="document not found")
    row = res.data[0]
    if not row.get("is_global") and row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this document")
    storage_path = row.get("pdf_storage_path")
    if not storage_path:
        raise HTTPException(status_code=404, detail="no PDF stored for this document")

    # storage_path is "legal-docs/<file>"; the storage SDK takes bucket + key
    bucket, _, key = storage_path.partition("/")
    try:
        signed = db.storage.from_(bucket).create_signed_url(key, expires_in)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="could not generate download link")

    return {
        "document_id": row["id"],
        "title": row.get("title"),
        "short_name": row.get("short_name"),
        "page_count": row.get("pdf_page_count"),
        "canonical_url": row.get("canonical_url"),
        "signed_url": signed.get("signedURL") or signed.get("signed_url") or signed.get("signedUrl"),
        "expires_in": expires_in,
    }


@router.post("/artifacts/sweep")
def sweep_old_artifacts(
    older_than_days: int = 30,
    dry_run: bool = True,
    x_admin_token: str | None = Header(default=None),
):
    """
    Soft-archive artifacts older than `older_than_days` (default 30) by
    stamping `archived_at`. With `dry_run=false`, also deletes the underlying
    storage objects. DESTRUCTIVE — gated behind the ADMIN_API_TOKEN secret
    (send it as the X-Admin-Token header). Intended for an external cron.
    """
    settings = get_settings()
    admin_token = getattr(settings, "admin_api_token", "") or os.environ.get("ADMIN_API_TOKEN", "")
    if not admin_token or x_admin_token != admin_token:
        raise HTTPException(status_code=403, detail="admin token required")
    from ..db.supabase import get_db
    from datetime import datetime, timedelta, timezone

    db = get_db()
    cutoff = (datetime.now(timezone.utc) - timedelta(days=older_than_days)).isoformat()

    # Candidates: rows older than cutoff that haven't been archived yet.
    rows = (
        db.table("artifacts")
        .select("id, storage_path, archived_at, created_at")
        .lt("created_at", cutoff)
        .is_("archived_at", "null")
        .limit(500)
        .execute()
    ).data or []

    archived = 0
    deleted_objects = 0
    errors: list[str] = []

    for row in rows:
        if dry_run:
            archived += 1
            continue
        # Delete the storage object first; if it fails we keep the row
        # un-stamped so a later sweep retries.
        path = row.get("storage_path") or ""
        if path and path.startswith("artifacts/"):
            key = path.split("/", 1)[1]
            try:
                db.storage.from_("artifacts").remove([key])
                deleted_objects += 1
            except Exception as e:  # noqa: BLE001
                errors.append(f"{row['id']}: {e}")
                continue
        try:
            db.table("artifacts").update(
                {"archived_at": datetime.now(timezone.utc).isoformat()}
            ).eq("id", row["id"]).execute()
            archived += 1
        except Exception as e:  # noqa: BLE001
            errors.append(f"{row['id']}: row-update {e}")

    return {
        "candidates": len(rows),
        "archived": archived,
        "deleted_storage_objects": deleted_objects,
        "errors": errors[:10],
        "cutoff": cutoff,
        "dry_run": dry_run,
    }


def _require_artifact_uuid(artifact_id: str) -> None:
    """Reject a malformed artifact id with a clean 404.

    artifacts.id is a uuid column, so a non-uuid path segment makes the
    Postgres cast raise and surface as a 500. Validate the format up front so
    a bad id reads as "not found" like every other missing artifact.
    """
    try:
        uuid.UUID(str(artifact_id))
    except (ValueError, TypeError, AttributeError):
        raise HTTPException(status_code=404, detail="artifact not found")


@router.get("/artifacts/{artifact_id}/pdf")
def get_artifact_pdf_url(artifact_id: str, expires_in: int = 3600, uid: str | None = Depends(optional_user)):
    """Signed URL for an agent-generated artifact PDF.

    Owned artifacts are only served to their owner; anonymous-demo artifacts
    (owner_id NULL, addressable only by their unguessable UUID) are served by
    capability.
    """
    _require_artifact_uuid(artifact_id)
    from ..db.supabase import get_db

    db = get_db()
    res = (
        db.table("artifacts")
        .select("id, title, kind, storage_path, page_count, size_bytes, source, meta, created_at, owner_id")
        .eq("id", artifact_id)
        .limit(1)
        .execute()
    )
    if not res.data:
        raise HTTPException(status_code=404, detail="artifact not found")
    row = res.data[0]
    if row.get("owner_id") and row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this artifact")
    storage_path = row.get("storage_path")
    if not storage_path or storage_path == "artifacts/pending":
        raise HTTPException(status_code=409, detail="artifact upload not finalized")

    bucket, _, key = storage_path.partition("/")
    try:
        signed = db.storage.from_(bucket).create_signed_url(key, expires_in)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="could not generate download link")

    return {
        "artifact_id": row["id"],
        "title": row.get("title"),
        "kind": row.get("kind"),
        "page_count": row.get("page_count"),
        "size_bytes": row.get("size_bytes"),
        "source": row.get("source"),
        "meta": row.get("meta"),
        "created_at": row.get("created_at"),
        "signed_url": signed.get("signedURL") or signed.get("signed_url") or signed.get("signedUrl"),
        "expires_in": expires_in,
    }


@router.get("/artifacts/{artifact_id}/text")
def get_artifact_text(artifact_id: str, uid: str | None = Depends(optional_user)):
    """Plain text of a generated document so the user can copy it directly.

    Mobile downloads of signed PDF URLs are unreliable, and users asked to just
    copy the drafted text. We return the stored source Markdown (same content
    that renders the PDF/Word), same authorization as the other artifact routes.
    """
    _require_artifact_uuid(artifact_id)
    from ..db.supabase import get_db

    db = get_db()
    res = (db.table("artifacts").select("id, title, meta, owner_id")
           .eq("id", artifact_id).limit(1).execute())
    if not res.data:
        raise HTTPException(status_code=404, detail="artifact not found")
    row = res.data[0]
    if row.get("owner_id") and row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this artifact")
    text = (row.get("meta") or {}).get("source_markdown")
    if not text:
        raise HTTPException(status_code=409, detail="this document has no copyable text")
    return {"artifact_id": artifact_id, "title": row.get("title"), "text": text}


@router.get("/artifacts/{artifact_id}/docx")
def get_artifact_docx_url(artifact_id: str, expires_in: int = 3600, uid: str | None = Depends(optional_user)):
    """Signed URL for the Word (.docx) version of an artifact, rendered on demand.

    Same authorization as the PDF route. The Word file is rendered from the
    artifact's stored source Markdown the first time it is requested, then
    cached in the bucket so subsequent downloads are instant.
    """
    _require_artifact_uuid(artifact_id)
    from ..db.supabase import get_db
    from ..services.docx_tools import ensure_artifact_docx

    db = get_db()
    res = (
        db.table("artifacts").select("id, title, meta, owner_id")
        .eq("id", artifact_id).limit(1).execute()
    )
    if not res.data:
        raise HTTPException(status_code=404, detail="artifact not found")
    row = res.data[0]
    if row.get("owner_id") and row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this artifact")

    try:
        out = ensure_artifact_docx(artifact_id)
    except ValueError as e:
        raise HTTPException(status_code=409, detail=str(e))
    except Exception:
        raise HTTPException(status_code=500, detail="could not render the Word document")

    bucket, _, key = out["storage_path"].partition("/")
    try:
        signed = db.storage.from_(bucket).create_signed_url(key, expires_in)
    except Exception:
        raise HTTPException(status_code=500, detail="could not generate download link")

    return {
        "artifact_id": artifact_id,
        "title": out.get("title"),
        "kind": "docx",
        "signed_url": signed.get("signedURL") or signed.get("signed_url") or signed.get("signedUrl"),
        "expires_in": expires_in,
    }


@router.get("/documents/by-title")
def get_document_by_title(title: str):
    """
    Look up a document by exact or fuzzy title match — used by the chat UI when
    a citation snapshot only knows the act name (not the document_id).
    """
    from ..db.supabase import get_db

    db = get_db()
    # Exact match first
    res = (
        db.table("legal_documents")
        .select("id, title, short_name, pdf_storage_path, pdf_page_count")
        .eq("title", title)
        .limit(1)
        .execute()
    )
    if not res.data:
        # Fall back to ILIKE so 'REPUBLIC OF ZAMBIA THE COMPANIES ACT' still
        # resolves to the new 'Companies Act, No. 10 of 2017' row.
        res = (
            db.table("legal_documents")
            .select("id, title, short_name, pdf_storage_path, pdf_page_count")
            .ilike("title", f"%{title.split()[-2] if len(title.split()) > 1 else title}%")
            .limit(5)
            .execute()
        )
    if not res.data:
        raise HTTPException(status_code=404, detail="no document matched")
    return {"matches": res.data}


@router.get("/documents")
def list_documents(
    session_id: str | None = None,
    folder_id: str | None = None,
    user_id: str | None = None,  # ignored — kept for backward-compat
    uid: str | None = Depends(optional_user),
):
    """
    List documents visible to the caller.

    - `global`   : the curated Zambian-law library (always available, even
                   to anonymous callers)
    - `owned`    : documents the AUTHENTICATED user uploaded (identity from
                   the bearer token, never the client-supplied user_id)
    - `attached` : documents attached to a chat session the user owns
    """
    from ..db.supabase import get_db

    db = get_db()

    cols = (
        "id, title, short_name, year, document_type, total_chunks, "
        "pdf_page_count, pdf_size_bytes, pdf_storage_path, canonical_url, "
        "is_global, owner_id, folder_id, created_at"
    )

    global_docs = (
        db.table("legal_documents").select(cols).eq("is_global", True)
        .order("title").execute().data or []
    )

    owned: list[dict] = []
    if uid:
        q = db.table("legal_documents").select(cols).eq("owner_id", uid)
        if folder_id == "unfiled":
            q = q.is_("folder_id", "null")
        elif folder_id:
            q = q.eq("folder_id", folder_id)
        owned = q.order("created_at", desc=True).execute().data or []

    attached: list[dict] = []
    if session_id and uid:
        # only expose attachments for a session the caller owns
        sess = db.table("chat_sessions").select("user_id").eq("id", session_id).limit(1).execute()
        if sess.data and sess.data[0].get("user_id") == uid:
            ids_res = (
                db.table("chat_session_documents").select("document_id")
                .eq("session_id", session_id).execute()
            )
            ids = [r["document_id"] for r in (ids_res.data or [])]
            if ids:
                attached = (
                    db.table("legal_documents").select(cols).in_("id", ids)
                    .order("title").execute().data or []
                )

    return {
        "global": global_docs,
        "owned": owned,
        "attached": attached,
        "counts": {
            "global": len(global_docs),
            "owned": len(owned),
            "attached": len(attached),
        },
    }


# ─── Folders (user-created groupings of uploaded documents) ──────────────────


class CreateFolderRequest(BaseModel):
    user_id: str
    name: str


class RenameFolderRequest(BaseModel):
    name: str


class MoveDocumentRequest(BaseModel):
    folder_id: str | None = None  # null = remove from any folder


@router.get("/folders")
def list_folders(user_id: str | None = None, uid: str = Depends(require_user)):
    """Return the authenticated user's folders + a per-folder document count."""
    from ..db.supabase import get_db
    db = get_db()
    folders = (
        db.table("document_folders").select("id, name, created_at")
        .eq("owner_id", uid).order("created_at", desc=False).execute().data or []
    )
    docs = (
        db.table("legal_documents").select("id, folder_id").eq("owner_id", uid)
        .execute().data or []
    )
    counts: dict[str, int] = {}
    unfiled = 0
    for d in docs:
        fid = d.get("folder_id")
        if fid is None:
            unfiled += 1
        else:
            counts[fid] = counts.get(fid, 0) + 1
    return {
        "folders": [
            {**f, "doc_count": counts.get(f["id"], 0)} for f in folders
        ],
        "unfiled_count": unfiled,
    }


@router.post("/folders")
def create_folder(request: CreateFolderRequest, uid: str = Depends(require_user)):
    from ..db.supabase import get_db
    db = get_db()
    name = request.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="name required")
    try:
        row = db.table("document_folders").insert(
            {"owner_id": uid, "name": name}
        ).execute()
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=409, detail="request could not be completed")
    return row.data[0] if row.data else {"status": "ok"}


@router.patch("/folders/{folder_id}")
def rename_folder(folder_id: str, request: RenameFolderRequest, uid: str = Depends(require_user)):
    from ..db.supabase import get_db
    name = request.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="name required")
    _assert_owns("document_folders", folder_id, uid)
    db = get_db()
    db.table("document_folders").update(
        {"name": name, "updated_at": "now()"}
    ).eq("id", folder_id).execute()
    return {"status": "ok"}


@router.delete("/folders/{folder_id}")
def delete_folder(folder_id: str, cascade: bool = False, uid: str = Depends(require_user)):
    """Delete a folder you own. Documents inside are unfiled (folder_id=null),
    not removed, unless cascade=true (destructive)."""
    from ..db.supabase import get_db
    _assert_owns("document_folders", folder_id, uid)
    db = get_db()
    if cascade:
        # only the caller's docs in that folder
        db.table("legal_documents").delete().eq("folder_id", folder_id).eq("owner_id", uid).execute()
    else:
        db.table("legal_documents").update({"folder_id": None}).eq("folder_id", folder_id).eq("owner_id", uid).execute()
    db.table("document_folders").delete().eq("id", folder_id).execute()
    return {"status": "ok", "cascade": cascade}


@router.patch("/documents/{document_id}/folder")
def move_document_to_folder(document_id: str, request: MoveDocumentRequest, uid: str = Depends(require_user)):
    """Move a document you own into a folder you own (or clear with null)."""
    from ..db.supabase import get_db
    _assert_owns("legal_documents", document_id, uid)
    if request.folder_id:
        _assert_owns("document_folders", request.folder_id, uid)
    db = get_db()
    db.table("legal_documents").update({"folder_id": request.folder_id}).eq("id", document_id).execute()
    return {"status": "ok", "folder_id": request.folder_id}


# ─── Per-thread document attachment ──────────────────────────────────────────


class AttachDocRequest(BaseModel):
    document_id: str


@router.get("/sessions/{session_id}/documents")
def list_session_documents(session_id: str, uid: str = Depends(require_user)):
    """Documents currently attached to a chat session the caller owns."""
    from ..db.supabase import get_db
    db = get_db()
    _assert_owns_session(session_id, uid)
    ids_res = (
        db.table("chat_session_documents").select("document_id, attached_at")
        .eq("session_id", session_id).execute()
    )
    rows = ids_res.data or []
    if not rows:
        return {"documents": []}
    ids = [r["document_id"] for r in rows]
    docs = (
        db.table("legal_documents")
        .select("id, title, short_name, total_chunks, pdf_page_count, owner_id, is_global")
        .in_("id", ids).execute().data or []
    )
    by_id = {d["id"]: d for d in docs}
    return {
        "documents": [
            {**by_id[r["document_id"]], "attached_at": r["attached_at"]}
            for r in rows if r["document_id"] in by_id
        ],
    }


@router.post("/sessions/{session_id}/documents/attach")
def attach_document(session_id: str, request: AttachDocRequest, uid: str = Depends(require_user)):
    """Attach a document to a chat session the caller owns."""
    from ..db.supabase import get_db
    db = get_db()
    _assert_owns_session(session_id, uid)
    # the document must be global or owned by the caller
    doc = db.table("legal_documents").select("owner_id,is_global").eq("id", request.document_id).limit(1).execute()
    if not doc.data:
        raise HTTPException(status_code=404, detail="document not found")
    d = doc.data[0]
    if not d.get("is_global") and d.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this document")
    try:
        db.table("chat_session_documents").upsert(
            {"session_id": session_id, "document_id": request.document_id},
        ).execute()
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=400, detail="request could not be completed")
    return {"status": "ok", "session_id": session_id, "document_id": request.document_id}


@router.delete("/sessions/{session_id}/documents/{document_id}")
def detach_document(session_id: str, document_id: str, uid: str = Depends(require_user)):
    """Remove a document attachment from a chat session the caller owns."""
    from ..db.supabase import get_db
    db = get_db()
    _assert_owns_session(session_id, uid)
    db.table("chat_session_documents").delete().eq("session_id", session_id).eq(
        "document_id", document_id
    ).execute()
    return {"status": "ok"}


INLINE_TIER_MAX_PAGES = 5


def _upload_object(db, bucket: str, key: str, content: bytes, content_type: str) -> None:
    """Upload bytes to a Supabase Storage object, replacing any existing object."""
    try:
        db.storage.from_(bucket).upload(
            path=key,
            file=content,
            file_options={"content-type": content_type, "upsert": "true"},
        )
    except Exception as e:
        # Older storage SDKs raise on duplicate even when upsert=true was sent;
        # fall back to remove + upload.
        if "exists" in str(e).lower() or "duplicate" in str(e).lower():
            try:
                db.storage.from_(bucket).remove([key])
            except Exception:
                pass
            db.storage.from_(bucket).upload(
                path=key,
                file=content,
                file_options={"content-type": content_type},
            )
        else:
            raise


def _extract_pdf_text(pdf_path: str) -> str:
    """Concatenate all pages' text. Used for the inline tier (small docs)."""
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            parts.append(t)
    return "\n\n".join(parts).strip()


@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    user_id: str | None = None,  # ignored — identity from token
    folder_id: str | None = None,
    uid: str = Depends(require_user),
):
    """Upload + ingest a PDF.

    Two paths, picked by page count:
      * **Inline tier** (≤5 pages) — extract full text, store it alongside the
        PDF in storage, mark the document with `total_chunks=0`. No embeddings.
        At chat time, the agent receives the full text in its system prompt.
      * **RAG tier** (>5 pages) — full chunk + embed via `ingest_pdf`.

    In both paths the PDF itself is uploaded to the `legal-docs` bucket so the
    citation viewer can render it.
    """
    import hashlib
    import io as _io

    from pypdf import PdfReader
    from ..db.supabase import get_db

    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    if folder_id:
        _assert_owns("document_folders", folder_id, uid)

    db = get_db()

    try:
        content = await file.read()
        size_bytes = len(content)
        pdf_hash = hashlib.sha256(content).hexdigest()

        # Dedup by hash — same bytes always map to the same document.
        existing = (
            db.table("legal_documents")
            .select("id,title,total_chunks,pdf_page_count")
            .eq("pdf_hash", pdf_hash)
            .execute()
            .data
        )
        if existing:
            doc = existing[0]
            is_inline = (doc.get("total_chunks", 0) or 0) == 0
            # Hitting the dedup path with an inline-tier doc means the user
            # has uploaded this exact file before — surface a promotion hint so
            # the UI can offer "Save to Library" (chunk + embed) once.
            return {
                "status": "skipped",
                "document_id": doc["id"],
                "chunks_created": doc.get("total_chunks", 0),
                "tier": "inline" if is_inline else "rag",
                "page_count": doc.get("pdf_page_count"),
                "suggest_promotion": is_inline,
            }

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            page_count = len(PdfReader(_io.BytesIO(content)).pages)
        except Exception:
            page_count = 0

        # ── Inline tier ────────────────────────────────────────────────────
        if 0 < page_count <= INLINE_TIER_MAX_PAGES:
            try:
                text = _extract_pdf_text(tmp_path)
            except Exception:
                text = ""

            title = file.filename.rsplit(".", 1)[0] or "Document"
            row = (
                db.table("legal_documents")
                .insert(
                    {
                        "title": title,
                        "short_name": title,
                        "document_type": "attachment",
                        "pdf_hash": pdf_hash,
                        "pdf_page_count": page_count,
                        "pdf_size_bytes": size_bytes,
                        "total_chunks": 0,
                        "total_sections": 0,
                        "is_global": False,
                        "owner_id": uid,
                        "folder_id": folder_id,
                    }
                )
                .execute()
                .data[0]
            )
            doc_id = row["id"]
            key_pdf = f"{doc_id}.pdf"
            try:
                _upload_object(db, "legal-docs", key_pdf, content, "application/pdf")
                if text:
                    _upload_object(
                        db,
                        "legal-docs",
                        f"{doc_id}.txt",
                        text.encode("utf-8"),
                        "text/plain; charset=utf-8",
                    )
                db.table("legal_documents").update(
                    {"pdf_storage_path": f"legal-docs/{key_pdf}"}
                ).eq("id", doc_id).execute()
            except Exception:
                pass

            os.unlink(tmp_path)
            return {
                "status": "inline",
                "document_id": doc_id,
                "chunks_created": 0,
                "tier": "inline",
                "page_count": page_count,
            }

        # ── RAG tier (>5 pages) ────────────────────────────────────────────
        result = ingest_pdf(tmp_path)
        doc = (result or {}).get("document") or {}
        doc_id = doc.get("id")
        if doc_id:
            patch: dict = {"is_global": False, "owner_id": uid}
            if folder_id:
                patch["folder_id"] = folder_id
            try:
                db.table("legal_documents").update(patch).eq("id", doc_id).execute()
            except Exception:
                pass
            # Upload PDF to storage so the citation viewer can render it.
            try:
                key_pdf = f"{doc_id}.pdf"
                _upload_object(db, "legal-docs", key_pdf, content, "application/pdf")
                db.table("legal_documents").update(
                    {"pdf_storage_path": f"legal-docs/{key_pdf}"}
                ).eq("id", doc_id).execute()
            except Exception:
                pass

        os.unlink(tmp_path)
        return {
            "status": result.get("status", "unknown"),
            "document_id": doc_id,
            "chunks_created": result.get("chunks_created", 0),
            "tier": "rag",
            "page_count": page_count,
        }
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=500, detail="request could not be completed")


@router.post("/documents/{doc_id}/promote")
def promote_document_to_library(
    doc_id: str,
    uid: str = Depends(require_user),
):
    """Promote an inline-tier document to full RAG (chunk + embed).

    Inline-tier docs have `total_chunks=0` and only their full text in storage.
    Promotion runs the standard parse → chunk → embed pipeline against the
    SAME row (no new document_id), so chips already pinned to user messages
    remain valid and the doc becomes searchable across all the user's chats.
    """
    from ..db.supabase import get_db
    from ..services.ingester import chunk_existing_pdf

    db = get_db()

    # Authorise: caller must own the doc (or admin). We use the same pattern
    # as _assert_owns but allow global docs to be promoted by anyone (they
    # already are searchable, so it's a no-op anyway).
    row_res = db.table("legal_documents").select(
        "id, owner_id, is_global, total_chunks, pdf_storage_path"
    ).eq("id", doc_id).limit(1).execute()
    if not row_res.data:
        raise HTTPException(status_code=404, detail="document not found")
    row = row_res.data[0]
    if not row.get("is_global") and row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this document")

    if (row.get("total_chunks") or 0) > 0:
        return {"status": "already_promoted", "document_id": doc_id, "chunks_created": row["total_chunks"]}

    storage_path = row.get("pdf_storage_path")
    if not storage_path:
        raise HTTPException(status_code=400, detail="document has no PDF in storage to promote")

    bucket, _, key = storage_path.partition("/")
    try:
        pdf_bytes = db.storage.from_(bucket).download(key)
    except Exception:
        raise HTTPException(status_code=502, detail="could not fetch PDF from storage")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    try:
        result = chunk_existing_pdf(tmp_path, doc_id)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    return {
        "status": result.get("status", "unknown"),
        "document_id": doc_id,
        "chunks_created": result.get("chunks_created", 0),
    }


@router.post("/brief/generate")
def generate_brief(request: BriefRequest):
    """
    Generate an IRAC legal brief from conversation history.

    Takes a list of conversation messages and produces a structured
    Issue, Rule, Application, Conclusion analysis based on the
    legal topics discussed.
    """
    if not request.messages:
        raise HTTPException(status_code=400, detail="Messages cannot be empty")

    try:
        user_message = build_irac_prompt(request.messages)

        result = generate_response(
            system_prompt=IRAC_SYSTEM_PROMPT,
            user_message=user_message,
            max_tokens=4096,
        )

        # Parse JSON from response
        answer = result["answer"]

        # Handle case where Claude wraps JSON in markdown code blocks
        if "```json" in answer:
            answer = answer.split("```json")[1].split("```")[0].strip()
        elif "```" in answer:
            answer = answer.split("```")[1].split("```")[0].strip()

        brief_data = json.loads(answer)
        return {
            "issue": brief_data.get("issue", ""),
            "rule": brief_data.get("rule", ""),
            "application": brief_data.get("application", ""),
            "conclusion": brief_data.get("conclusion", ""),
            "citations": brief_data.get("citations", []),
        }
    except json.JSONDecodeError:
        # If JSON parsing fails, return the raw text in the issue field
        return {
            "issue": answer,
            "rule": "",
            "application": "",
            "conclusion": "",
            "citations": [],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail="request could not be completed")


class BriefExportRequest(BaseModel):
    issue: str = ""
    rule: str = ""
    application: str = ""
    conclusion: str = ""
    citations: list[dict] = []
    format: str = "pdf"  # "pdf" | "docx"
    title: str = "Legal Brief"


_IRAC_SECTIONS = [
    ("issue", "Issue"),
    ("rule", "Rule"),
    ("application", "Application"),
    ("conclusion", "Conclusion"),
]


def _brief_filename(title: str, ext: str) -> str:
    base = re.sub(r"[^A-Za-z0-9]+", "-", title or "legal-brief").strip("-").lower()
    return f"{base or 'legal-brief'}.{ext}"


def _citation_line(c: dict) -> str:
    act = (c.get("act") or "").strip()
    section = (c.get("section") or "").strip()
    page = c.get("page")
    parts = ", ".join(p for p in [act, section] if p)
    if isinstance(page, int) and page > 0:
        parts = f"{parts} (p.{page})" if parts else f"p.{page}"
    return parts


@router.post("/brief/export")
def export_brief(request: BriefExportRequest):
    """Render an IRAC brief into a downloadable PDF or DOCX file."""
    fmt = (request.format or "pdf").lower()
    if fmt not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="format must be 'pdf' or 'docx'")

    if not any(getattr(request, key) for key, _ in _IRAC_SECTIONS):
        raise HTTPException(status_code=400, detail="brief has no content to export")

    subtitle = "IRAC Analysis — " + datetime.now(timezone.utc).strftime("%d %B %Y")
    citations = [c for c in request.citations if _citation_line(c)]

    try:
        if fmt == "pdf":
            from ..services.pdf_tools import _render_markdown_pdf

            parts: list[str] = []
            for key, label in _IRAC_SECTIONS:
                value = (getattr(request, key) or "").strip()
                if value:
                    parts.append(f"## {label}\n\n{value}")
            if citations:
                lines = "\n".join(f"- {_citation_line(c)}" for c in citations)
                parts.append(f"## Sources Referenced\n\n{lines}")
            body_md = "\n\n".join(parts)

            pdf_bytes = _render_markdown_pdf(
                title=request.title or "Legal Brief",
                body_md=body_md,
                subtitle=subtitle,
            )
            if not pdf_bytes:
                raise HTTPException(status_code=500, detail="failed to render PDF")
            return StreamingResponse(
                io.BytesIO(pdf_bytes),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'attachment; filename="{_brief_filename(request.title, "pdf")}"'
                },
            )

        # docx
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        heading = doc.add_heading(request.title or "Legal Brief", level=0)
        sub = doc.add_paragraph(subtitle)
        sub.runs[0].italic = True
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for key, label in _IRAC_SECTIONS:
            value = (getattr(request, key) or "").strip()
            if not value:
                continue
            doc.add_heading(label, level=1)
            doc.add_paragraph(value)

        if citations:
            doc.add_heading("Sources Referenced", level=1)
            for c in citations:
                doc.add_paragraph(_citation_line(c), style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{_brief_filename(request.title, "docx")}"'
            },
        )
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=500, detail="request could not be completed")


# ─── Templates ───────────────────────────────────────────────────────────────


class UpdateTemplateRequest(BaseModel):
    name: str | None = None
    description: str | None = None


class CreateTemplateFolderRequest(BaseModel):
    user_id: str
    name: str


class RenameTemplateFolderRequest(BaseModel):
    name: str


class MoveTemplateRequest(BaseModel):
    folder_id: str | None = None  # null clears the folder


@router.get("/templates")
def list_templates(user_id: str | None = None, folder_id: str | None = None, uid: str = Depends(require_user)):
    """List the authenticated user's uploaded templates, optionally by folder."""
    from ..services.templates import list_templates_for_owner

    rows = list_templates_for_owner(uid, folder_id=folder_id)
    return {"templates": rows, "count": len(rows)}


@router.get("/template-folders")
def list_template_folders(user_id: str | None = None, uid: str = Depends(require_user)):
    """Return the authenticated user's template folders + counts."""
    from ..db.supabase import get_db

    db = get_db()
    folders = (
        db.table("template_folders")
        .select("id, name, created_at")
        .eq("owner_id", uid)
        .order("created_at", desc=False)
        .execute()
        .data
        or []
    )
    rows = (
        db.table("templates")
        .select("id, folder_id")
        .eq("owner_id", uid)
        .execute()
        .data
        or []
    )
    counts: dict[str, int] = {}
    unfiled = 0
    for r in rows:
        fid = r.get("folder_id")
        if fid is None:
            unfiled += 1
        else:
            counts[fid] = counts.get(fid, 0) + 1
    return {
        "folders": [{**f, "doc_count": counts.get(f["id"], 0)} for f in folders],
        "unfiled_count": unfiled,
    }


@router.post("/template-folders")
def create_template_folder(req: CreateTemplateFolderRequest, uid: str = Depends(require_user)):
    from ..db.supabase import get_db

    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="name required")
    db = get_db()
    try:
        row = (
            db.table("template_folders")
            .insert({"owner_id": uid, "name": name})
            .execute()
        )
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=409, detail="request could not be completed")
    return row.data[0] if row.data else {"status": "ok"}


@router.patch("/template-folders/{folder_id}")
def rename_template_folder(folder_id: str, req: RenameTemplateFolderRequest, uid: str = Depends(require_user)):
    from ..db.supabase import get_db

    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="name required")
    _assert_owns("template_folders", folder_id, uid)
    db = get_db()
    db.table("template_folders").update(
        {"name": name, "updated_at": "now()"}
    ).eq("id", folder_id).execute()
    return {"status": "ok"}


@router.delete("/template-folders/{folder_id}")
def delete_template_folder(folder_id: str, cascade: bool = False, uid: str = Depends(require_user)):
    """Delete a template folder you own. By default templates inside become
    unfiled; with cascade=true they're deleted (storage included)."""
    from ..db.supabase import get_db

    _assert_owns("template_folders", folder_id, uid)
    db = get_db()
    if cascade:
        # Pull storage paths first so we can clean blobs.
        rows = (
            db.table("templates")
            .select("id, storage_path")
            .eq("folder_id", folder_id)
            .execute()
            .data
            or []
        )
        for row in rows:
            path = row.get("storage_path") or ""
            if path:
                bucket, _, key = path.partition("/")
                try:
                    db.storage.from_(bucket).remove([key])
                except Exception:
                    pass
        db.table("templates").delete().eq("folder_id", folder_id).execute()
    else:
        db.table("templates").update({"folder_id": None}).eq("folder_id", folder_id).execute()
    db.table("template_folders").delete().eq("id", folder_id).execute()
    return {"status": "ok", "cascade": cascade}


@router.patch("/templates/{template_id}/folder")
def move_template_to_folder(template_id: str, req: MoveTemplateRequest, uid: str = Depends(require_user)):
    from ..db.supabase import get_db

    _assert_owns("templates", template_id, uid)
    if req.folder_id:
        _assert_owns("template_folders", req.folder_id, uid)
    db = get_db()
    db.table("templates").update({"folder_id": req.folder_id}).eq(
        "id", template_id
    ).execute()
    return {"status": "ok", "folder_id": req.folder_id}


@router.post("/templates/upload")
async def upload_template(
    file: UploadFile = File(...),
    user_id: str | None = None,  # ignored — identity from token
    name: str | None = None,
    description: str | None = None,
    folder_id: str | None = None,
    uid: str = Depends(require_user),
):
    """
    Upload a template file (.docx | .pdf | .txt | .md).

    Stored in the private `templates` Supabase Storage bucket; a row in the
    `templates` table records owner + name + description + a preview of the
    text content for the agent's `suggest_templates` tool.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="file required")
    if folder_id:
        _assert_owns("template_folders", folder_id, uid)

    from ..services.templates import (
        TEMPLATES_BUCKET,
        extract_preview,
        file_type_for,
        upload_to_storage,
    )
    from ..db.supabase import get_db

    file_type = file_type_for(file.filename)
    if not file_type:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Use .docx, .pdf, .txt, or .md.",
        )

    content = await file.read()
    size = len(content)
    preview, page_count = extract_preview(content, file_type)

    # Default name = filename without extension
    template_name = (name or os.path.splitext(file.filename)[0]).strip()
    if not template_name:
        raise HTTPException(status_code=400, detail="name required")

    try:
        storage_path = upload_to_storage(content, file_type, uid)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="upload failed")

    db = get_db()
    try:
        row = (
            db.table("templates")
            .insert(
                {
                    "owner_id": uid,
                    "name": template_name,
                    "description": (description or "").strip() or None,
                    "file_type": file_type,
                    "file_size_bytes": size,
                    "storage_path": storage_path,
                    "preview_text": preview or None,
                    "page_count": page_count,
                    "folder_id": folder_id,
                }
            )
            .execute()
        ).data
    except Exception as e:  # noqa: BLE001
        # Roll back the storage write so we don't leak orphan blobs.
        try:
            bucket, _, key = storage_path.partition("/")
            db.storage.from_(bucket).remove([key])
        except Exception:
            pass
        raise HTTPException(status_code=409, detail="request could not be completed")

    return {"template": row[0] if row else None}


@router.patch("/templates/{template_id}")
def update_template(template_id: str, req: UpdateTemplateRequest, uid: str = Depends(require_user)):
    """Rename a template you own or update its description."""
    from ..db.supabase import get_db

    _assert_owns("templates", template_id, uid)
    patch: dict = {"updated_at": "now()"}
    if req.name is not None:
        name = req.name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="name cannot be empty")
        patch["name"] = name
    if req.description is not None:
        patch["description"] = req.description.strip() or None

    db = get_db()
    db.table("templates").update(patch).eq("id", template_id).execute()
    return {"status": "ok"}


@router.delete("/templates/{template_id}")
def delete_template(template_id: str, uid: str = Depends(require_user)):
    """Delete a template you own (DB row + storage object)."""
    from ..db.supabase import get_db

    db = get_db()
    res = (
        db.table("templates")
        .select("id, storage_path, owner_id")
        .eq("id", template_id)
        .limit(1)
        .execute()
    )
    if not res.data:
        raise HTTPException(status_code=404, detail="template not found")
    row = res.data[0]
    if row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this template")
    path = row.get("storage_path") or ""
    if path:
        bucket, _, key = path.partition("/")
        try:
            db.storage.from_(bucket).remove([key])
        except Exception:
            # Storage delete is best-effort — never block row deletion.
            pass
    db.table("templates").delete().eq("id", template_id).execute()
    return {"status": "ok"}


@router.get("/templates/{template_id}/file")
def get_template_signed_url(template_id: str, expires_in: int = 3600, uid: str = Depends(require_user)):
    """Short-lived signed URL to download/preview a template you own."""
    from ..db.supabase import get_db

    db = get_db()
    res = (
        db.table("templates")
        .select("id, name, file_type, storage_path, page_count, file_size_bytes, owner_id")
        .eq("id", template_id)
        .limit(1)
        .execute()
    )
    if not res.data:
        raise HTTPException(status_code=404, detail="template not found")
    row = res.data[0]
    if row.get("owner_id") != uid:
        raise HTTPException(status_code=403, detail="not authorized for this template")
    path = row.get("storage_path") or ""
    if not path:
        raise HTTPException(status_code=404, detail="template has no stored file")
    bucket, _, key = path.partition("/")
    try:
        signed = db.storage.from_(bucket).create_signed_url(key, expires_in)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="could not generate download link")
    return {
        "template_id": row["id"],
        "name": row.get("name"),
        "file_type": row.get("file_type"),
        "page_count": row.get("page_count"),
        "size_bytes": row.get("file_size_bytes"),
        "signed_url": signed.get("signedURL")
        or signed.get("signed_url")
        or signed.get("signedUrl"),
        "expires_in": expires_in,
    }
